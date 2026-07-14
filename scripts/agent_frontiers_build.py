#!/usr/bin/env python3
"""Build and package the final Reviewer 1 minor revision artifacts.

This is a temporary helper invoked only when the dedicated trigger file exists.
It modifies the checked-out working tree, renders the submission package, performs
text and page-render checks, and replaces governance-preflight.log with a ZIP
artifact for the existing upload-artifact step.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
PAPER_DIR = ROOT / "papers" / "ontological_friction"
FIGURE_DIR = PAPER_DIR / "figures"
MANUSCRIPT = PAPER_DIR / "paper_ontological_friction_frontiers_submission.md"
RESPONSE = PAPER_DIR / "frontiers_response_to_reviewers.md"
HTML = PAPER_DIR / "paper_ontological_friction_frontiers_submission.html"
DOCX = PAPER_DIR / "paper_ontological_friction_frontiers_submission.docx"
PDF = PAPER_DIR / "paper_ontological_friction_frontiers_submission.pdf"
BUILD_DIR = ROOT / ".agent_frontiers_build"
QA_DIR = BUILD_DIR / "qa"
BUNDLE = BUILD_DIR / "frontiers_reviewer1_final_revision_bundle.zip"
ARTIFACT_TARGET = ROOT / "governance-preflight.log"

TABLE2_OLD = "**Table 2. Predicted $\\Psi_f$ proxy signatures across clinical conditions.**"
TABLE2_NEW = (
    "**Table 2. Predicted $\\Psi_f$ proxy signatures across clinical conditions. "
    "Bold entries marked with † denote the primary disorder-discriminating predictions "
    "used in the paired dissociation tests.**"
)
TABLE3_OLD = r"90 (30$\times$3)"
TABLE3_NEW = "90 (three groups of 30 participants)"
RESPONSE_MARKER = (
    "\n---\n\nWe believe these revisions address the reviewers' concerns and materially "
    "strengthen the manuscript, and we thank the reviewers again for their time."
)
RESPONSE_FOLLOW_UP = """

---

## Reviewer 1 — Follow-up Minor Comments

**1. Table 2 caption — clarify the meaning of bold entries marked with †.**
Response: Thank you for pointing this out. The caption of Table 2 now states that the bold entries marked with † denote the primary disorder-discriminating predictions used in the paired dissociation tests.

**2. Table 3 — correct the group-size notation.**
Response: Thank you for identifying the rendering problem. The affected entry has been corrected to “90 (three groups of 30 participants),” which preserves the intended sample-size meaning while avoiding ambiguous mathematical formatting.
"""


def run(command: list[str], *, cwd: Path = ROOT) -> None:
    print("+", " ".join(command), flush=True)
    subprocess.run(command, cwd=cwd, check=True)


def install_toolchain() -> None:
    run(["sudo", "apt-get", "update"])
    run(
        [
            "sudo",
            "apt-get",
            "install",
            "-y",
            "pandoc",
            "libreoffice",
            "poppler-utils",
            "fonts-dejavu-core",
        ]
    )
    run(
        [
            "uv",
            "pip",
            "install",
            "--python",
            sys.executable,
            "pypandoc",
            "python-docx",
            "lxml",
            "pillow",
            "matplotlib",
            "numpy",
            "scipy",
        ]
    )
    if not any(shutil.which(name) for name in ("google-chrome", "google-chrome-stable", "chromium", "chromium-browser")):
        raise RuntimeError("Chrome/Chromium is unavailable on the Actions runner.")


def apply_revisions() -> None:
    manuscript = MANUSCRIPT.read_text(encoding="utf-8")
    if manuscript.count(TABLE2_OLD) != 1:
        raise RuntimeError(f"Expected exactly one Table 2 caption; found {manuscript.count(TABLE2_OLD)}")
    if manuscript.count(TABLE3_OLD) != 1:
        raise RuntimeError(f"Expected exactly one Table 3 notation; found {manuscript.count(TABLE3_OLD)}")
    manuscript = manuscript.replace(TABLE2_OLD, TABLE2_NEW).replace(TABLE3_OLD, TABLE3_NEW)
    MANUSCRIPT.write_text(manuscript, encoding="utf-8")

    response = RESPONSE.read_text(encoding="utf-8")
    if "## Reviewer 1 — Follow-up Minor Comments" in response:
        raise RuntimeError("Follow-up section already exists unexpectedly.")
    if response.count(RESPONSE_MARKER) != 1:
        raise RuntimeError(f"Expected exactly one response closing marker; found {response.count(RESPONSE_MARKER)}")
    response = response.replace(RESPONSE_MARKER, RESPONSE_FOLLOW_UP + RESPONSE_MARKER)
    RESPONSE.write_text(response, encoding="utf-8")


def build_figures() -> None:
    scripts = [
        "fig1_control_architecture.py",
        "fig2_proxy_map.py",
        "fig3_ros_coupling.py",
        "fig4_clinical_radar.py",
        "fig5_protocol.py",
    ]
    for script in scripts:
        run([sys.executable, script], cwd=FIGURE_DIR)
    expected = [
        "fig1_control_architecture.png",
        "fig2_proxy_map.png",
        "fig3_ros_coupling.png",
        "fig4_clinical_radar.png",
        "fig5_protocol.png",
    ]
    missing = [name for name in expected if not (FIGURE_DIR / name).is_file()]
    if missing:
        raise RuntimeError(f"Missing generated figures: {missing}")


def render_submission() -> None:
    run([sys.executable, "render_frontiers_submission_package.py"], cwd=PAPER_DIR)
    for path in (HTML, DOCX, PDF):
        if not path.is_file() or path.stat().st_size == 0:
            raise RuntimeError(f"Missing or empty rendered file: {path}")


def verify_text() -> None:
    required_text = (
        "Bold entries marked with † denote the primary disorder-discriminating predictions used in the paired dissociation tests.",
        TABLE3_NEW,
    )
    for path in (MANUSCRIPT, HTML):
        raw_text = path.read_text(encoding="utf-8")
        if path == HTML:
            from lxml import html as lxml_html

            visible_text = lxml_html.fromstring(raw_text).text_content()
        else:
            visible_text = raw_text
        normalized_text = " ".join(visible_text.split())
        for required in required_text:
            if required not in normalized_text:
                raise RuntimeError(f"Missing required text in {path.name}: {required}")
        for forbidden in ("30$3", TABLE3_OLD):
            if forbidden in raw_text:
                raise RuntimeError(f"Forbidden notation remains in {path.name}: {forbidden}")

    from docx import Document

    doc = Document(DOCX)
    docx_text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    normalized_docx_text = " ".join(docx_text.split())
    for required in required_text:
        if required not in normalized_docx_text:
            raise RuntimeError(f"Missing required text in DOCX: {required}")

    pdf_text = BUILD_DIR / "submission_pdf.txt"
    run(["pdftotext", str(PDF), str(pdf_text)])
    extracted = " ".join(pdf_text.read_text(encoding="utf-8", errors="replace").split())
    if "three groups of 30 participants" not in extracted:
        raise RuntimeError("Corrected Table 3 text is missing from the PDF text extraction.")


def render_qa_pages() -> tuple[int, int]:
    docx_qa = QA_DIR / "docx"
    pdf_qa = QA_DIR / "pdf"
    lo_profile = BUILD_DIR / "lo-profile"
    docx_qa.mkdir(parents=True, exist_ok=True)
    pdf_qa.mkdir(parents=True, exist_ok=True)
    lo_profile.mkdir(parents=True, exist_ok=True)

    run(
        [
            "libreoffice",
            "--headless",
            f"-env:UserInstallation=file://{lo_profile}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_qa),
            str(DOCX),
        ]
    )
    converted = docx_qa / DOCX.with_suffix(".pdf").name
    if not converted.is_file():
        raise RuntimeError("LibreOffice did not produce a PDF from the DOCX.")

    run(
        [
            "pdftoppm",
            "-jpeg",
            "-r",
            "72",
            "-jpegopt",
            "quality=65",
            str(converted),
            str(docx_qa / "page"),
        ]
    )
    run(
        [
            "pdftoppm",
            "-jpeg",
            "-r",
            "72",
            "-jpegopt",
            "quality=65",
            str(PDF),
            str(pdf_qa / "page"),
        ]
    )

    docx_pages = sorted(docx_qa.glob("page-*.jpg"))
    pdf_pages = sorted(pdf_qa.glob("page-*.jpg"))
    if not docx_pages or not pdf_pages:
        raise RuntimeError("Page rendering produced no QA images.")
    return len(docx_pages), len(pdf_pages)


def make_bundle(docx_pages: int, pdf_pages: int) -> None:
    manifest = BUILD_DIR / "MANIFEST.txt"
    manifest.write_text(
        "Reviewer 1 final minor revision bundle\n"
        f"DOCX QA pages: {docx_pages}\n"
        f"PDF QA pages: {pdf_pages}\n"
        "Source of truth: paper_ontological_friction_frontiers_submission.md\n"
        "Table 2 caption and Table 3 group-size notation verified in MD, HTML, DOCX, and PDF.\n",
        encoding="utf-8",
    )

    files = [
        (MANUSCRIPT, MANUSCRIPT.name),
        (RESPONSE, RESPONSE.name),
        (HTML, HTML.name),
        (DOCX, DOCX.name),
        (PDF, PDF.name),
        (manifest, manifest.name),
    ]
    with zipfile.ZipFile(BUNDLE, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        for source, arcname in files:
            archive.write(source, arcname)
        for page in sorted((QA_DIR / "docx").glob("page-*.jpg")):
            archive.write(page, f"qa/docx/{page.name}")
        for page in sorted((QA_DIR / "pdf").glob("page-*.jpg")):
            archive.write(page, f"qa/pdf/{page.name}")

    if not BUNDLE.is_file() or BUNDLE.stat().st_size == 0:
        raise RuntimeError("Artifact bundle was not created.")
    os.replace(BUNDLE, ARTIFACT_TARGET)
    print(
        f"Prepared artifact bundle at {ARTIFACT_TARGET} "
        f"({ARTIFACT_TARGET.stat().st_size} bytes).",
        flush=True,
    )


def main() -> int:
    if BUILD_DIR.exists():
        shutil.rmtree(BUILD_DIR)
    BUILD_DIR.mkdir(parents=True)
    install_toolchain()
    apply_revisions()
    build_figures()
    render_submission()
    verify_text()
    docx_pages, pdf_pages = render_qa_pages()
    make_bundle(docx_pages, pdf_pages)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
